import io
import json
import os
from pathlib import Path
from typing import Any, Dict, Tuple

from anthropic import Anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request, send_file
from fpdf import FPDF


SYSTEM_PROMPT = """
You are an expert resume writer with 15 years of experience crafting ATS-optimized,
visually polished resumes for technical roles.

Your goal is to tailor the candidate's resume to the provided job description while
matching a strict visual format and fitting within a 1-page length constraint.

═══════════════════════════════════════
CANDIDATE WORK AUTHORIZATION STATUS
═══════════════════════════════════════
The candidate is on an F1 Student Visa (OPT/STEM OPT).
This means:
  ✅ Can work in the US without employer visa sponsorship
  ✅ Eligible for roles that say "no sponsorship required" or
     "must be authorized to work in the US"
  ❌ NOT a US Citizen
  ❌ NOT a Green Card / Permanent Resident holder
  ❌ NOT eligible for roles requiring security clearance

Use this status to drive the eligibility logic in STEP 0 below.

═══════════════════════════════════════
INPUTS
═══════════════════════════════════════
CANDIDATE RESUME:
{resume}

TARGET JOB DESCRIPTION:
{job_description}

═══════════════════════════════════════
STEP 0 — ELIGIBILITY CHECK (RUN FIRST)
═══════════════════════════════════════
Before doing ANYTHING else, scan the full job description carefully and
classify it into one of three buckets:

─────────────────────────────────────
BUCKET A — HARD BLOCK ❌
─────────────────────────────────────
Trigger if the JD contains ANY of these (exact or semantic match):
  - "US Citizenship required"
  - "Must be a US Citizen"
  - "Green Card required"
  - "Permanent Resident required"
  - "Only US Citizens and Green Card holders eligible"
  - "Active or ability to obtain Security Clearance required"
  - "Must hold or be eligible for [government] clearance"

ACTION: STOP. Do NOT generate a resume.
Return ONLY this output:

---ELIGIBILITY CHECK FAILED---

⚠️ This job is NOT suitable for your profile.

Reason: This position requires [exact phrase from JD — e.g., "US Citizenship"
/ "Green Card" / "Security Clearance"], which is restricted to US Citizens
or Permanent Residents only.

Your Status: F1 Visa (OPT/STEM OPT) — you do not meet this requirement
regardless of sponsorship.

Recommendation: Skip this application and look for roles that state:
  ✅ "Open to all work authorizations"
  ✅ "Will sponsor H-1B"
  ✅ "F1/OPT candidates welcome"
  ✅ No citizenship or clearance restrictions mentioned

─────────────────────────────────────
BUCKET B — ELIGIBLE ✅
─────────────────────────────────────
Trigger if the JD contains ANY of these:
  - "Must be authorized to work in the US"
  - "No visa sponsorship available" / "We do not sponsor visas"
  - "Must be able to work without sponsorship"
  - "Employment authorization required"
  - No mention of citizenship, Green Card, or clearance at all

ACTION: Candidate qualifies — F1 OPT/STEM OPT satisfies work
authorization without requiring sponsorship.
Proceed silently to Section 1. Do NOT mention the eligibility
check in the resume output.

─────────────────────────────────────
BUCKET C — AMBIGUOUS ⚠️
─────────────────────────────────────
Trigger if the JD has unclear or mixed signals, such as:
  - "Preferred: US Citizen or Green Card" (preferred, not required)
  - "Security clearance a plus"
  - Sponsorship language is vague or contradictory

ACTION: Do NOT block. Generate the resume AND add this note at the top:

---ELIGIBILITY NOTE---

⚠️ Advisory: This job description contains unclear work authorization
language: "[exact phrase from JD]"

Your Status: F1 Visa (OPT/STEM OPT) — you CAN work without sponsorship
but are not a US Citizen or Green Card holder.

Recommendation: Proceed with the application but confirm directly with
the recruiter whether F1 OPT candidates are considered before investing
time in interviews.

[Resume follows below]

═══════════════════════════════════════
SECTION 1 — PROFILE SUMMARY (CRITICAL)
═══════════════════════════════════════
Write a 4-sentence profile summary under the PROFILE section.
Follow this EXACT sentence-by-sentence structure:

SENTENCE 1 — WHO YOU ARE:
  Formula: [Degree] + [Years of Experience] + [Top 2–3 domain areas from the JD]
  Purpose: Establishes identity and seniority in the first line.
  Rule:    Extract domain areas directly from the job description — use their
           exact language, not synonyms.
  Example: "Master's in Data Science graduate with 2 years of experience in
            data analysis, machine learning, and business intelligence."

SENTENCE 2 — WHAT YOU KNOW (Tools & Tech):
  Formula: [Proficiency statement] + [Tools/Languages] + [Platforms/Frameworks from JD]
  Purpose: This is your ATS sentence — pack it with JD-matching keywords.
  Rule:    Only include tools that appear in BOTH the candidate resume AND the JD.
           Prioritize JD tools over resume tools when trimming.
  Example: "Proficient in Python, SQL, Tableau, and Power BI, with hands-on
            exposure to Big Data technologies including Hadoop, Spark, and Hive."

SENTENCE 3 — WHAT YOU'VE DONE (Skills in Action):
  Formula: [Experience areas] + [Specialized competencies] + [Technical focus areas]
  Purpose: Bridges tools → real application. Shows depth beyond just knowing the tools.
  Rule:    Use action-oriented phrases. Reflect the responsibilities listed in the JD.
           Do NOT repeat tools already mentioned in Sentence 2.
  Example: "Experienced in predictive modelling, statistical analysis, ETL pipeline
            development, and cloud-based analytics with a focus on anomaly detection
            and data visualization."

SENTENCE 4 — WHY YOU DO IT (Passion + Value):
  Formula: [Passion/mission statement] + [Aligned with company's goal or role impact]
  Purpose: Humanizes the resume, signals cultural fit, and closes with energy.
  Rule:    Tailor this to the company's mission or the role's stated impact if
           mentioned in the JD. Avoid generic phrases like "hardworking" or
           "passionate learner" without context.
  Example: "Passionate about leveraging AI/ML to derive actionable business insights
            and solve real-world challenges at scale."

PROFILE SUMMARY STYLE RULES (APPLY TO ALL 4 SENTENCES):
  - Write in dense paragraph format — NO bullet points in this section
  - NO first-person pronouns ("I", "my", "me") — use third-person implied tone
  - Bold ALL technical terms, tools, domain phrases, and quantifiers
    (e.g., **Python**, **SQL**, **2 years of experience**, **machine learning**)
  - Keep total length to 3–5 lines maximum — recruiters scan, not read
  - Mirror the job description's exact language wherever possible
  - NO fabrication — only use skills and experience from the candidate's resume

═══════════════════════════════════════
SECTION 2 — VISUAL FORMATTING RULES
═══════════════════════════════════════
1. MARGINS & SPACING:
   - Content must be dense to fit a Narrow Margin (0.5 inch) layout.
   - NO empty lines between section text, horizontal separators,
     and the following section header.

2. CONTACT HEADER (CENTERED):
   - Candidate name centered at the top in larger bold font.
   - Directly below (single centered line):
     Location | Email | Phone | LinkedIn | GitHub
   - Insert a horizontal line "---" immediately after contact info.

3. SECTION HEADERS:
   Use these EXACT titles in ALL CAPS and BOLD with no extra spaces or colons:
   **PROFILE**
   **EXPERIENCE**
   **EDUCATION**
   **LEADERSHIP**
   **TECHNICAL SKILLS**
   **ACADEMIC PROJECTS**
   **CERTIFICATIONS**
   **HOBBIES**

4. HORIZONTAL SEPARATORS:
   - Insert "---" immediately after the contact header and after
     every section's content block.
   - No blank lines before or after the separator.

5. EXPERIENCE & LEADERSHIP STRUCTURE:
   - Line 1: **Organization – Location** (Bold)
   - Line 2: **Role | Start Date – End Date** (Bold)
   - Followed by compact bullet points using "-"
   - Bold all key technical terms and metrics within bullets
     (e.g., **Python**, **SQL**, **30%**, **95%**, **20+ clients**)

6. ACADEMIC PROJECTS:
   - Every project title MUST be in **BOLD ALL CAPS**
     (e.g., **GPT FROM SCRATCH**, **SALES FORECASTING DASHBOARD**)

═══════════════════════════════════════
SECTION 3 — CONTENT & LENGTH RULES
═══════════════════════════════════════
- MAX LENGTH: Strictly 1 page. Prune aggressively to fit.
- SMART PRUNING: Remove the least relevant academic projects or leadership
  bullets first. Prioritize the 2–3 projects most relevant to the JD.
- NO FABRICATION: Only use information present in the provided resume.
- PRESERVE THESE METRICS EXACTLY (do not alter or omit):
    - 4.0 GPA
    - 30% query optimization improvement
    - 95% report accuracy improvement

═══════════════════════════════════════
SECTION 4 — OUTPUT FORMAT
═══════════════════════════════════════
Return your response in this exact structure:

---RESUME---
[Full tailored resume following all formatting rules above]

---KEYWORDS MATCHED---
[Bulleted list of JD keywords found and used in the resume]

---ATS SCORE---
Score: [X/100]

Breakdown:
- Keyword Match Rate (40 pts): [score] — [brief explanation]
- Formatting & Readability (30 pts): [score] — [brief explanation]
- Experience Relevance (30 pts): [score] — [brief explanation]

Overall Verdict: [1–2 sentence summary of resume strength for this JD]
"""


MODEL_NAME = "claude-sonnet-4-6"
ALLOWED_FONT_SIZES = {8, 10, 12}
INPUT_TOKEN_PRICE = 0.000001
OUTPUT_TOKEN_PRICE = 0.000005

BASE_DIR = Path(__file__).resolve().parent
DATA_FILE = BASE_DIR / "data.json"

app = Flask(__name__)
load_dotenv(BASE_DIR / ".env", override=True)


def _default_settings() -> Dict[str, Any]:
    return {"base_resume": "", "font_size": 10}


def normalize_font_size(value: Any) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return 10
    return parsed if parsed in ALLOWED_FONT_SIZES else 10


def load_settings() -> Dict[str, Any]:
    default = _default_settings()
    if not DATA_FILE.exists():
        return default

    try:
        with DATA_FILE.open("r", encoding="utf-8") as file:
            loaded = json.load(file)
    except (json.JSONDecodeError, OSError):
        return default

    base_resume = str(loaded.get("base_resume", ""))
    font_size = normalize_font_size(loaded.get("font_size", 10))
    return {"base_resume": base_resume, "font_size": font_size}


def save_settings(settings: Dict[str, Any]) -> None:
    serializable = {
        "base_resume": str(settings.get("base_resume", "")),
        "font_size": normalize_font_size(settings.get("font_size", 10)),
    }
    with DATA_FILE.open("w", encoding="utf-8") as file:
        json.dump(serializable, file, indent=2)


def parse_model_output(text: str) -> Dict[str, str]:
    result = {"resume": "", "keywords": "", "ats_score": ""}
    cleaned = text.strip()

    sections = [
        ("---RESUME---", "resume"),
        ("---KEYWORDS MATCHED---", "keywords"),
        ("---ATS SCORE---", "ats_score"),
    ]

    for idx, (header, key) in enumerate(sections):
        if header not in cleaned:
            continue
        start = cleaned.find(header) + len(header)
        if idx < len(sections) - 1:
            next_header = sections[idx + 1][0]
            end = cleaned.find(next_header, start)
            chunk = cleaned[start:end if end != -1 else None]
        else:
            chunk = cleaned[start:]
        result[key] = chunk.strip()

    if not result["resume"]:
        result["resume"] = cleaned

    return result


def compute_cost(usage: Any) -> Dict[str, Any]:
    input_tokens = 0
    output_tokens = 0

    if usage is not None:
        input_tokens = int(getattr(usage, "input_tokens", 0) or 0)
        output_tokens = int(getattr(usage, "output_tokens", 0) or 0)

    input_cost = input_tokens * INPUT_TOKEN_PRICE
    output_cost = output_tokens * OUTPUT_TOKEN_PRICE
    total_cost = input_cost + output_cost

    return {
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "input_cost": round(input_cost, 8),
        "output_cost": round(output_cost, 8),
        "total_cost": round(total_cost, 8),
        "display": f"${total_cost:.8f}",
    }


def get_anthropic_client() -> Anthropic:
    api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("Missing ANTHROPIC_API_KEY environment variable.")
    return Anthropic(api_key=api_key)


def build_messages(base_resume: str, job_description: str) -> list[Dict[str, Any]]:
    return [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": f"BASE_RESUME:\n{base_resume}",
                    "cache_control": {"type": "ephemeral"},
                },
                {
                    "type": "text",
                    "text": f"JOB_DESCRIPTION:\n{job_description}",
                },
            ],
        }
    ]


def sanitize_pdf_text(text: str) -> str:
    replacements = {
        "\u2022": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\xa0": " ",
    }
    normalized = text
    for source, target in replacements.items():
        normalized = normalized.replace(source, target)
    return normalized.encode("latin-1", errors="replace").decode("latin-1")


def split_markdown_bold_segments(text: str) -> list[tuple[str, bool]]:
    segments: list[tuple[str, bool]] = []
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part:
            continue
        segments.append((part, idx % 2 == 1))
    if not segments:
        return [("", False)]
    return segments


def generate_pdf_bytes(resume_text: str, font_size: int) -> bytes:
    pdf = FPDF()
    narrow_margin_mm = 12.7
    pdf.set_margins(narrow_margin_mm, narrow_margin_mm, narrow_margin_mm)
    pdf.set_auto_page_break(auto=True, margin=narrow_margin_mm)
    pdf.add_page()
    pdf.set_font("Times", size=font_size)

    line_height = max(3.2, font_size * 0.42)
    rendered_line_count = 0
    for line in resume_text.splitlines():
        stripped = sanitize_pdf_text(line.strip())
        if not stripped:
            continue
        plain_for_center = stripped.replace("**", "")
        if rendered_line_count < 2:
            pdf.set_font("Times", style="B", size=font_size)
            pdf.cell(0, 10, plain_for_center, align="C", ln=1)
            pdf.set_font("Times", size=font_size)
            rendered_line_count += 1
            continue
        if stripped.replace("_", "") == "" and len(stripped) >= 10:
            y_pos = min(pdf.get_y() + 2, pdf.h - pdf.b_margin)
            pdf.set_y(y_pos)
            pdf.line(pdf.l_margin, y_pos, pdf.w - pdf.r_margin, y_pos)
            pdf.set_y(y_pos + 2)
            pdf.set_x(pdf.l_margin)
            rendered_line_count += 1
            continue
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, line_height, stripped, markdown=True)
        rendered_line_count += 1

    payload = pdf.output(dest="S")
    if isinstance(payload, (bytes, bytearray)):
        return bytes(payload)
    return payload.encode("latin-1", errors="replace")


def generate_docx_bytes(resume_text: str, font_size: int) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(font_size)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    rendered_line_count = 0
    for line in resume_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.replace("_", "") == "" and len(stripped) >= 10:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0

            p_pr = paragraph._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "0")
            bottom.set(qn("w:color"), "auto")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)

            rendered_line_count += 1
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        if rendered_line_count < 2:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for segment_text, segment_bold in split_markdown_bold_segments(stripped):
            run = paragraph.add_run(segment_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(font_size)
            run.bold = segment_bold or (rendered_line_count < 2)
        rendered_line_count += 1

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/settings", methods=["GET"])
def get_settings():
    return jsonify(load_settings())


@app.route("/settings", methods=["POST"])
def update_settings():
    data = request.get_json(silent=True) or {}
    base_resume = str(data.get("base_resume", "")).strip()
    font_size = normalize_font_size(data.get("font_size", 10))

    payload = {"base_resume": base_resume, "font_size": font_size}
    save_settings(payload)
    return jsonify(payload)


@app.route("/tailor", methods=["POST"])
def tailor_resume():
    data = request.get_json(silent=True) or {}
    base_resume = str(data.get("base_resume", "")).strip()
    job_description = str(data.get("job_description", "")).strip()

    if not base_resume:
        return jsonify({"error": "Base resume cannot be empty."}), 400
    if not job_description:
        return jsonify({"error": "Job description cannot be empty."}), 400

    try:
        client = get_anthropic_client()
        response = client.messages.create(
            model=MODEL_NAME,
            max_tokens=1800,
            system=SYSTEM_PROMPT,
            messages=build_messages(base_resume, job_description),
        )
        response_text = "".join(
            block.text for block in getattr(response, "content", []) if block.type == "text"
        ).strip()
        parsed = parse_model_output(response_text)
        costs = compute_cost(getattr(response, "usage_metadata", None))
        if costs["input_tokens"] == 0 and costs["output_tokens"] == 0:
            costs = compute_cost(getattr(response, "usage", None))

        return jsonify(
            {
                "resume": parsed["resume"],
                "keywords_matched": parsed["keywords"],
                "ats_score": parsed["ats_score"],
                "raw_response": response_text,
                "cost": costs,
            }
        )
    except Exception as exc:
        error_message = str(exc).strip() or "Unknown API error."
        if "timeout" in error_message.lower():
            return (
                jsonify({"error": "Anthropic API timeout. Please try again."}),
                500,
            )
        return jsonify({"error": f"Failed to tailor resume: {error_message}"}), 500


def _validate_download_input(data: Dict[str, Any]) -> Tuple[str, int, Any]:
    resume_text = str(data.get("resume_text", "")).strip()
    font_size = normalize_font_size(data.get("font_size", 10))
    if not resume_text:
        return "", font_size, (jsonify({"error": "Resume text is required."}), 400)
    return resume_text, font_size, None


@app.route("/download/pdf", methods=["POST"])
def download_pdf():
    data = request.get_json(silent=True) or {}
    resume_text, font_size, error_response = _validate_download_input(data)
    if error_response:
        return error_response

    try:
        pdf_bytes = generate_pdf_bytes(resume_text, font_size)
        return send_file(
            io.BytesIO(pdf_bytes),
            as_attachment=True,
            download_name="tailored_resume.pdf",
            mimetype="application/pdf",
        )
    except Exception as exc:
        return jsonify({"error": f"Failed to generate PDF: {str(exc).strip()}"}), 500


@app.route("/download/docx", methods=["POST"])
def download_docx():
    data = request.get_json(silent=True) or {}
    resume_text, font_size, error_response = _validate_download_input(data)
    if error_response:
        return error_response

    docx_bytes = generate_docx_bytes(resume_text, font_size)
    return send_file(
        io.BytesIO(docx_bytes),
        as_attachment=True,
        download_name="tailored_resume.docx",
        mimetype=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )


if __name__ == "__main__":
    if not DATA_FILE.exists():
        save_settings(_default_settings())
    app.run(debug=True)
